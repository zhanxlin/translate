

import os
from docx import Document

patient="Banswer + symptom"
doctor="Yrequest+symptom"

path="F:/2018/大四/对话系统/数据集/胃溃疡/胃溃疡/胃溃疡/"
files_list=os.listdir(path)

print(files_list)





def rewrite(file_read,file_write):
    f = open(path+file_write+".txt", "w")
    document=Document(path+file_read)
    size=len(document.paragraphs)
    print(size)
    for i,line in enumerate(document.paragraphs):
        if (i>1) & (i!=size-1):
            print(i,size-1,(i>1)&(i!=size-1))
            if i%2==1:
                print(line.text+patient)
                f.write(line.text+patient+"\n")
            elif i%2==0:
                print(line.text+doctor)
                f.write(line.text+doctor+"\n")


for i,each in enumerate(files_list):
    print(each,str(i))
    rewrite(each,str(i))




